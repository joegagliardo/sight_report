import sys, os, io, json, re, datetime
from typing import Any, Optional, List, Dict
from pydantic import BaseModel, Field
import vertexai
from vertexai.generative_models import Part, Image
from google.cloud import storage
from fpdf import FPDF
from PIL import Image, ImageDraw, ImageFont, ImageOps
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from google.auth import default
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload

from pyairtable import Api

def process_gcs_manifest_tool(input_json: str) -> List[Any]:
    """
    Parses a manifest of names and GCS paths, and returns multimodal parts for Gemini.
    """
    data: Dict[str, str] = json.loads(input_json)
    multimodal_parts: List[Any] = []
    multimodal_parts.append("Please process the following images. The first is a template, followed by individuals:")

    for key, gcs_uri in data.items():
        mime_type = "image/png" if gcs_uri.lower().endswith(".png") else "image/jpeg"
        image_part = Part.from_uri(uri=gcs_uri, mime_type=mime_type)
        multimodal_parts.append(f"Label for next image: {key}")
        multimodal_parts.append(image_part)

    multimodal_parts.append("Analyze these images and return the result.")
    return multimodal_parts


def generate_trip_infographic(data):
    print('calling generate_trip_infographic')
    
    # Initialize variables to avoid UnboundLocalError
    attendee_logo = None
    storage_client = None
    
    # --- Configuration & Constants ---
    width = 1200
    row_height = 400
    header_height = 250
    bg_color = "#0D2B63"  # Dark Navy
    bright_blue = "#4DA6FF"
    white = "#FFFFFF"
    black = "#000000"
    
    classes = data.get("classes", [])
    num_classes = len(classes)
    
    # Calculate rows based on custom mapping for 1-15 items
    layout_map = {
        1: [1], 2: [2], 3: [3], 4: [4], 5: [5],
        6: [3, 3], 7: [4, 3], 8: [4, 4], 9: [5, 4], 10: [5, 5],
        11: [4, 4, 3], 12: [4, 4, 4], 13: [5, 4, 4], 14: [5, 5, 4], 15: [5, 5, 5]
    }
    
    distribution = layout_map.get(num_classes, [5] * (num_classes // 5) + ([num_classes % 5] if num_classes % 5 else []))
    max_across = max(distribution) if distribution else 0
    center_layout = max_across < 5
    
    rows = []
    current_idx = 0
    for count in distribution:
        rows.append(classes[current_idx : current_idx + count])
        current_idx += count
    
    total_height = header_height + (len(rows) * row_height) + 100
    img = Image.new('RGB', (width, total_height), color=white)
    draw = ImageDraw.Draw(img)

    font_path = "fonts/Inter.ttf"
    font_bold_path = "fonts/Inter_24pt-Bold.ttf"
    try:
        if not os.path.exists(font_path):
            raise FileNotFoundError(f"Font file not found at {font_path}")

        font_bold = ImageFont.truetype(font_bold_path, 24)
        font_bold2 = ImageFont.truetype(font_bold_path, 32)
        font_bold3 = ImageFont.truetype(font_bold_path, 15)

        font_reg = ImageFont.truetype(font_path, 15)
        font_large = ImageFont.truetype(font_path, 40)
        
    except Exception as e:
        print(f"Warning: Could not load Inter.ttf ({e}). Falling back to default.")
        font_bold = font_reg = font_large = ImageFont.load_default()
        font_bold2 = font_bold3 = font_bold

    try:
        storage_client = storage.Client()
    except Exception as e:
        print(f"Warning: Could not initialize storage client ({e})")

    # --- 1. Header Construction ---
    # Fetch & Draw Google Cloud Logo from GCS
    try:
        if not storage_client:
            raise Exception("Storage client not initialized")
        bucket = storage_client.bucket("roitraining-dashboard-grounding")
        
        # Try to fetch Cloud TRIP Report logo
        try:
            gc_logo_blob = bucket.blob("company_logos/Cloud TRIP Report logo.png")
            gc_logo_bytes = gc_logo_blob.download_as_bytes()
            gc_logo = Image.open(io.BytesIO(gc_logo_bytes))
            
            # Trim whitespace from Google Cloud logo
            if gc_logo.mode != 'RGBA':
                gc_logo = gc_logo.convert('RGBA')
            bbox = gc_logo.getbbox()
            if bbox:
                gc_logo = gc_logo.crop(bbox)
                
            # Scale Google Cloud logo
            gc_logo_h = 80
            aspect = gc_logo.width / gc_logo.height
            gc_logo = gc_logo.resize((int(gc_logo_h * aspect), gc_logo_h), Image.Resampling.LANCZOS)
            
            img.paste(gc_logo, (50, 42), mask=gc_logo if gc_logo.mode == 'RGBA' else None)
            title_x = 50 + int(gc_logo_h * aspect) + 15
        except Exception as e:
            print(f"Warning: Could not fetch Google Cloud logo ({e})")
            title_x = 50

        # Try to fetch attendee icon
        try:
            attendee_logo_blob = bucket.blob("company_logos/attendee icon.png")
            attendee_bytes = attendee_logo_blob.download_as_bytes()
            attendee_logo = Image.open(io.BytesIO(attendee_bytes)).convert("RGBA")
            
            # Scale to fit nicely
            icon_h = 27
            aspect = attendee_logo.width / attendee_logo.height
            attendee_logo = attendee_logo.resize((int(icon_h * aspect), icon_h), Image.Resampling.LANCZOS)
        except Exception as e:
            print(f"Warning: Could not fetch attendee logo ({e})")

    except Exception as e:
        print(f"Warning: Header construction error ({e})")
        title_x = 50

    # --- Fetch & Draw Company Logo from GCS ---
    try:
        if not storage_client:
            raise Exception("Storage client not initialized")
        bucket_name = "roitraining-dashboard-grounding"
        blob_name = f"company_logos/{data['company']}.png"
        
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        
        logo_bytes = blob.download_as_bytes()
        company_logo = Image.open(io.BytesIO(logo_bytes))
        
        if company_logo.mode != 'RGBA':
            company_logo = company_logo.convert('RGBA')
        bbox = company_logo.getbbox()
        if bbox:
            company_logo = company_logo.crop(bbox)
            
        max_logo_h = 56
        aspect_ratio = company_logo.width / company_logo.height
        new_w = int(max_logo_h * aspect_ratio)
        company_logo = company_logo.resize((new_w, max_logo_h), Image.Resampling.LANCZOS)
        
        company_logo_y = 122 + 10
        img.paste(company_logo, (100, company_logo_y), mask=company_logo if company_logo.mode == 'RGBA' else None)
        print(f"Successfully added logo: {blob_name}")
    except Exception as e:
        print(f"Warning: Could not fetch company logo from GCS ({e}). Leaving blank.")

    # --- 2. Dark Blue Section ---
    draw.rectangle([0, header_height - 50, width, total_height], fill=bg_color)
    
    title_text = "Completed PLLJ Training Overview"
    company_text = data["company"]
    
    if center_layout:
        bbox_title = draw.textbbox((0, 0), title_text, font=font_bold)
        title_w = bbox_title[2] - bbox_title[0]
        draw.text(((width - title_w) // 2, header_height - 30), title_text, fill=white, font=font_bold)
        
        bbox_company = draw.textbbox((0, 0), company_text.upper(), font=font_bold)
        company_w = bbox_company[2] - bbox_company[0]
        draw.text(((width - company_w) // 2, header_height + 14), company_text.upper(), fill=bright_blue, font=font_bold2)
    else:
        draw.text((50, header_height - 30), title_text, fill=white, font=font_large)
        draw.text((50, header_height + 14), company_text.upper(), fill=bright_blue, font=font_bold2)

    # --- 3. Timeline Rows ---
    current_y = header_height + 150
    margin = 50
    items_per_row = 5
    min_gap = 20
    
    content_width = width - (2 * margin)
    card_w = (content_width - (items_per_row - 1) * min_gap) / items_per_row
    
    padding = margin + (card_w / 2)
    fixed_step = card_w + min_gap

    instructor_cache = {}

    for row in rows:
        num_in_row = len(row)
        if center_layout:
            row_width = (num_in_row - 1) * fixed_step + card_w
            row_start_x = (width - row_width) / 2
            x_positions = [int(row_start_x + card_w/2 + (i * fixed_step)) for i in range(num_in_row)]
        else:
            x_positions = [int(padding + (i * fixed_step)) for i in range(num_in_row)]

        if num_in_row > 1:
            draw.line([(x_positions[0], current_y), (x_positions[-1], current_y)], fill=bright_blue, width=3)

        for i, entry in enumerate(row):
            x = x_positions[i]
            instructor_name = entry.get('instructor', 'Unknown')
            
            bar_half_length = 40
            line_top = current_y - bar_half_length
            line_bottom = current_y + bar_half_length
            
            draw.line([(x, line_top), (x, line_bottom)], fill=bright_blue, width=3)

            pill_w, pill_h = 140, 50
            pill_top = line_top - pill_h
            draw.rounded_rectangle([x - pill_w//2, pill_top, x + pill_w//2, line_top], 
                                   radius=20, fill="#345491")
            
            draw.text((x - 45, pill_top + 15), entry.get('date', 'Unknown'), fill=white, font=font_reg)

            card_h = 180
            card_top = current_y + 40
            draw.rounded_rectangle([x - card_w//2, card_top, x + card_w//2, card_top + card_h], 
                                   radius=15, fill="#2A4A8E")
            
            title = entry.get('title', 'No Title')
            max_title_w = card_w - 20
            words = title.split()
            title_lines = []
            while words and len(title_lines) < 3:
                line = ""
                while words and draw.textbbox((0, 0), line + words[0], font=font_reg)[2] < max_title_w:
                    line += (words.pop(0) + " ")
                if not line: line = words.pop(0)
                title_lines.append(line.strip())
            
            if words:
                title_lines[-1] = title_lines[-1][:max(0, len(title_lines[-1])-3)] + "..."
            
            for j, line in enumerate(title_lines):
                draw.text((x - card_w//2 + 10, card_top + 10 + (j * 20)), line, fill=white, font=font_reg)
            
            attendee_margin = 10
            photo_size = 50
            attendee_text = str(entry.get('attendees', 0))
            
            if attendee_logo:
                at_icon_x = x - card_w//2 + attendee_margin + (photo_size // 2) - (attendee_logo.width // 2)
                at_icon_y = card_top + card_h - attendee_logo.height - attendee_margin
                
                at_text_x = at_icon_x + attendee_logo.width + 8
                at_text_y = at_icon_y + 6
                
                img.paste(attendee_logo, (int(at_icon_x), int(at_icon_y)), mask=attendee_logo)
                draw.text((at_text_x, at_text_y), attendee_text, fill=white, font=font_bold3)
                ref_y = at_icon_y
            else:
                emoji_text = f"👤 {attendee_text}"
                ebbox = draw.textbbox((0, 0), emoji_text, font=font_bold3)
                emoji_width = ebbox[2] - ebbox[0]
                at_text_x = x - card_w//2 + attendee_margin + (photo_size // 2) - (emoji_width // 2)
                at_text_y = card_top + card_h - 25 - attendee_margin
                draw.text((at_text_x, at_text_y), emoji_text, fill=white, font=font_bold3)
                ref_y = at_text_y

            photo_x = x - card_w//2 + attendee_margin
            photo_y = ref_y - photo_size - 10
            circle_bbox = [photo_x, photo_y, photo_x + photo_size, photo_y + photo_size]
            
            try:
                if instructor_name not in instructor_cache:
                    blob_name = f"instructor_photos/{instructor_name}.jpg"
                    bucket = storage_client.bucket("roitraining-dashboard-grounding")
                    blob = bucket.blob(blob_name)
                    photo_bytes = blob.download_as_bytes()
                    photo_img = Image.open(io.BytesIO(photo_bytes)).convert("RGBA")
                    crop_res = 100
                    photo_img = ImageOps.fit(photo_img, (crop_res, crop_res), centering=(0.5, 0.5))
                    mask = Image.new('L', (crop_res, crop_res), 0)
                    mask_draw = ImageDraw.Draw(mask)
                    mask_draw.ellipse((0, 0, crop_res, crop_res), fill=255)
                    circular_photo = Image.new('RGBA', (crop_res, crop_res), (0, 0, 0, 0))
                    circular_photo.paste(photo_img, (0, 0), mask=mask)
                    circular_photo = circular_photo.resize((photo_size, photo_size), Image.Resampling.LANCZOS)
                    instructor_cache[instructor_name] = circular_photo
                
                img.paste(instructor_cache[instructor_name], (int(circle_bbox[0]), int(circle_bbox[1])), 
                          mask=instructor_cache[instructor_name])
            except Exception as e:
                print(f"Warning: Could not fetch instructor photo for {instructor_name} ({e}). Leaving blank.")

            name_x = photo_x + photo_size + 10
            name_parts = (instructor_name.split(' ') + [""])[:2]
            draw.text((name_x, photo_y + 5), name_parts[0], fill=white, font=font_reg)
            draw.text((name_x, photo_y + 25), name_parts[1], fill=white, font=font_reg)

            dot_radius = 6
            draw.ellipse([x - dot_radius, line_top, x + dot_radius, line_top + 2 * dot_radius], fill=bright_blue)
            draw.ellipse([x - dot_radius, line_bottom - 2 * dot_radius, x + dot_radius, line_bottom], fill=bright_blue)

        current_y += row_height

    # --- 4. Footer ---
    try:
        roi_logo_blob = bucket.blob("company_logos/ROI.png")
        roi_logo_bytes = roi_logo_blob.download_as_bytes()
        roi_logo = Image.open(io.BytesIO(roi_logo_bytes))
        
        if roi_logo.mode != 'RGBA':
            roi_logo = roi_logo.convert('RGBA')
        roi_bbox = roi_logo.getbbox()
        if roi_bbox:
            roi_logo = roi_logo.crop(roi_bbox)
        
        roi_h = 40
        aspect = roi_logo.width / roi_logo.height
        roi_logo = roi_logo.resize((int(roi_h * aspect), roi_h), Image.Resampling.LANCZOS)
        
        img.paste(roi_logo, (50, total_height - 65), mask=roi_logo if roi_logo.mode == 'RGBA' else None)
    except Exception as e:
        print(f"Warning: Could not fetch ROI logo for footer ({e})")
        draw.text((50, total_height - 60), "🌐 ROI Training", fill=white, font=font_reg)

    # Save Output
    date_str = datetime.datetime.now().strftime("%Y%m%d")
    clean_company = data["company"].replace(" ", "_")
    output_path = f"reports/{clean_company}_infographic_{date_str}.png"
    img.save(output_path)
    
    # Try to save to bucket if possible
    # try:
    #     # Save to the 'reports' folder in GCS using only the filename
    #     gcs_dest = f"reports/{os.path.basename(output_path)}"
    #     save_to_bucket(output_path, 'roitraining-dashboard-grounding', gcs_dest)
    # except Exception as e:
    #     print(f"Warning: Could not upload infographic to bucket ({e})")
    
    print(f"Infographic generated: {output_path}")
    
    return {
        "status": "success",
        "infographic_path": output_path,
        "filename": output_path
    }


def save_report_as_pdf(company_name: str, report_text: str, infographic_path: str = "last_infographic.png") -> str:
    """
    Generates a PDF report containing the infographic and analysis text.
    Layout: Page 1 = Infographic, Page 2+ = Analysis text.
    Returns the local path of the generated PDF.
    """
    try:
        pdf = FPDF()
        
        # 1. Page 1: Infographic
        pdf.add_page()
        if os.path.exists(infographic_path):
            pdf.image(infographic_path, x=10, y=10, w=190)
        else:
            pdf.set_font("helvetica", 'B', 16)
            pdf.cell(0, 10, f"TRIP Report - {company_name}", ln=True, align='C')
            pdf.ln(10)

        # 2. Page 2 and subsequent: Report Text
        pdf.add_page()
        pdf.set_font("helvetica", size=12)
        
        sanitized_text = report_text.replace("’", "'").replace("“", '"').replace("”", '"').replace("–", "-")
        
        lines = sanitized_text.split('\n')
        for line in lines:
            if not line.strip():
                pdf.ln(5)
                continue
            if line.startswith('# '):
                pdf.set_font("helvetica", 'B', 14)
                pdf.cell(0, 10, line[2:], ln=True)
                pdf.set_font("helvetica", size=12)
            elif line.startswith('## '):
                pdf.set_font("helvetica", 'B', 13)
                pdf.cell(0, 10, line[3:], ln=True)
                pdf.set_font("helvetica", size=12)
            elif line.startswith('### '):
                pdf.set_font("helvetica", 'B', 12)
                pdf.cell(0, 10, line[4:], ln=True)
                pdf.set_font("helvetica", size=12)
            else:
                pdf.multi_cell(190, 7, line)
        
        date_str = datetime.datetime.now().strftime("%Y%m%d")
        clean_company = company_name.replace(" ", "_")
        local_pdf = f"reports/{clean_company}_TRIP_Report_{date_str}.pdf"
        pdf.output(local_pdf)
        
        print(f"Successfully generated PDF: {local_pdf}")
        return local_pdf
        
    except Exception as e:
        error_msg = f"Error generating PDF: {str(e)}"
        print(error_msg)
        return error_msg

def save_to_bucket(local_file_path: str, bucket_name: str = "roitraining-dashboard-grounding", destination_path: str = None) -> str:
    """
    Saves a local file to a designated storage bucket.
    """
    try:
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        
        if not destination_path:
            destination_path = f"reports/{os.path.basename(local_file_path)}"
            
        blob = bucket.blob(destination_path)
        blob.upload_from_filename(local_file_path)
        
        gcs_uri = f"gs://{bucket_name}/{destination_path}"
        print(f"Successfully uploaded {local_file_path} to {gcs_uri}")
        
        return gcs_uri
    except Exception as e:
        error_msg = f"Error uploading to bucket: {str(e)}"
        print(error_msg)
        return error_msg


class DocImageObject(BaseModel):
    infographic_path: Optional[str] = Field(None, description="Local path to the generated infographic png")
    filename: Optional[str] = Field(None, description="Filename of the infographic")
    status: Optional[str] = Field(None, description="Status of the generation process")

def create_and_share_google_doc(company_name: str, report_text: str, gcs_image_uri: Optional[str] = None, share_email: str = "TRIP.Reports@roitraining.com", folder_id: Optional[str] = None, local_image_path: Optional[str] = None, image_object: Optional[Dict[str, Any]] = None) -> str:
    """
    Prioritizes local_image_path or image_object by uploading/locating it and generating a reachable URL.
    """
    # 0. Resolve folder_id from environment if not provided (prevents LLM corruption of long IDs)
    folder_id = folder_id or os.environ.get("DRIVE_FOLDER_ID")

    try:
        # 0. Extract from image_object if provided
        if image_object:
            print(f"DEBUG: Processing image_object: {image_object}")
            # image_object is passed as a dict by the ADK
            obj_path = image_object.get("infographic_path") or image_object.get("filename")
                
            if not local_image_path:
                local_image_path = obj_path
            
            if not gcs_image_uri:
                gcs_image_uri = image_object.get("gcs_uri")
        
        scopes = [
            'https://www.googleapis.com/auth/documents', 
            'https://www.googleapis.com/auth/drive.file',
            'https://www.googleapis.com/auth/drive'
        ]
        
        creds = None
        # Use hardcoded path if running locally (no K_SERVICE env var), otherwise use environment
        if not os.environ.get("K_SERVICE"):
            creds_path = "/Users/joey/Dev/sight_report/service-account-key.json"
        else:
            creds_path = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")
            
        if creds_path and os.path.exists(creds_path):
            print(f"DEBUG: Using explicit service account file: {creds_path}")
            creds = service_account.Credentials.from_service_account_file(creds_path, scopes=scopes)
        else:
            print("DEBUG: GOOGLE_APPLICATION_CREDENTIALS not found or invalid. Falling back to default credentials.")
            creds, _ = default(scopes=scopes)
            if hasattr(creds, 'with_scopes'):
                creds = creds.with_scopes(scopes)

        # Log the identity being used
        if hasattr(creds, 'service_account_email'):
            print(f"DEBUG: Authenticated as Service Account: {creds.service_account_email}")
        else:
            # For ADC, we might not have the email directly in the creds object
            print(f"DEBUG: Authenticated using ADC (Type: {type(creds)})")

        docs_service = build('docs', 'v1', credentials=creds)
        drive_service = build('drive', 'v3', credentials=creds)
        
        clean_company = company_name.replace(" ", "_")
        date_str = datetime.datetime.now().strftime('%Y%m%d')
        title = f"{clean_company}_TRIP_Report_{date_str}"
        
        # We use the Drive API to create the file because it is already verified to work for this Service Account, 
        # whereas the Docs API create() sometimes fails with obscure 403s.
        doc_metadata = {
            'name': title,
            'mimeType': 'application/vnd.google-apps.document'
        }
        if folder_id:
            doc_metadata['parents'] = [folder_id]
            
        file = drive_service.files().create(
            body=doc_metadata, 
            fields='id',
            supportsAllDrives=True
        ).execute()
        doc_id = file.get('id')
        print(f"Created Google Doc {doc_id} using Drive API" + (f" in folder {folder_id}" if folder_id else ""))

        requests = []
        
        # 1. Parse Markdown links and prepare clean text
        links = []
        clean_text = ""
        last_idx = 0
        # Regex for [text](url)
        for match in re.finditer(r'\[(.*?)\]\((.*?)\)', report_text):
            # Text before the link
            clean_text += report_text[last_idx:match.start()]
            
            # Start offset in clean_text
            start_offset = len(clean_text) + 1
            link_text = match.group(1)
            link_url = match.group(2)
            
            clean_text += link_text
            # End offset
            end_offset = len(clean_text) + 1
            
            links.append({
                'start': start_offset,
                'end': end_offset,
                'url': link_url
            })
            last_idx = match.end()
        
        clean_text += report_text[last_idx:]

        # 2. Insert the clean text body
        requests.append({
            'insertText': {
                'location': {'index': 1},
                'text': clean_text
            }
        })

        # 3. Add styling for each link
        for link in links:
            requests.append({
                'updateTextStyle': {
                    'range': {
                        'startIndex': link['start'],
                        'endIndex': link['end']
                    },
                    'textStyle': {
                        'link': {'url': link['url']},
                        'foregroundColor': {
                            'color': {
                                'rgbColor': {'blue': 0.8, 'green': 0.3, 'red': 0.1}
                            }
                        },
                        'underline': True
                    },
                    'fields': 'link,foregroundColor,underline'
                }
            })
        
        public_url = None
        
        # Priority: Local Image Path (Find existing upload and Sign)
        if local_image_path:
            try:
                # We know generate_trip_infographic already uploaded the file to gs://{bucket}/reports/{filename}
                # So we just find it and sign it to avoid 403 upload errors.
                filename = os.path.basename(local_image_path)
                bucket_name = os.environ.get("LOGO_BUCKET", "roitraining-dashboard-grounding")
                blob_name = f"reports/{filename}"
                
                print(f"DEBUG: Locating existing upload for Google Doc: gs://{bucket_name}/{blob_name}")
                storage_client = storage.Client(credentials=creds, project=os.environ.get("PROJECT_ID", "roitraining-dashboard"))
                bucket = storage_client.bucket(bucket_name)
                blob = bucket.blob(blob_name)
                
                # Generate Signed URL (This is a local calculation, no API call needed if we have the key)
                try:
                    public_url = blob.generate_signed_url(
                        version="v4",
                        expiration=datetime.timedelta(minutes=15),
                        method="GET"
                    )
                except Exception as sign_e:
                    print(f"DEBUG: V4 Signing failed ({sign_e}), attempting fallback to public URL...")
                    # FALLBACK: Make the blob public for the duration of the request
                    # This is often needed on Cloud Run where signBlob permission is missing.
                    try:
                        blob.make_public()
                        public_url = blob.public_url
                        print(f"DEBUG: Made blob public: {public_url}")
                    except Exception as pub_e:
                        print(f"Warning: Could not make blob public ({pub_e}).")
                        public_url = f"https://storage.googleapis.com/{bucket_name}/{blob_name}"

                print(f"DEBUG: Reachable URL for Google Doc: {public_url[:50]}...")
            except Exception as e:
                print(f"Warning: Failed to locate or sign existing upload ({e}). Falling back to GCS URI if provided.")

        # Fallback: GCS URI (Sign)
        if not public_url and gcs_image_uri:
            if gcs_image_uri.startswith("gs://"):
                try:
                    parts = gcs_image_uri[5:].split("/", 1)
                    bucket_name = parts[0]
                    blob_name = parts[1]
                    
                    storage_client = storage.Client(credentials=creds, project=os.environ.get("PROJECT_ID", "roitraining-dashboard"))
                    bucket = storage_client.bucket(bucket_name)
                    blob = bucket.blob(blob_name)
                    
                    public_url = blob.generate_signed_url(
                        version="v4",
                        expiration=datetime.timedelta(minutes=15),
                        method="GET"
                    )
                    print(f"DEBUG: Generated Signed URL for GCS image: {public_url[:50]}...")
                except Exception as sign_e:
                    print(f"Warning: Could not sign GCS URI ({sign_e}). Falling back to public link.")
                    public_url = f"https://storage.googleapis.com/{bucket_name}/{blob_name}"
            else:
                public_url = gcs_image_uri

        # Insert Image and Text in reverse order of appearance (since we insert at index 1)
        # Final Order: [Image] -> [\n\n] -> [report_text]
        
        # 1. Insert Image (Must be last request for index 1 to be at the very top)
        if public_url:
            # We add the newline spacing FIRST so the image can take index 1
            requests.append({
                'insertText': {
                    'location': {'index': 1},
                    'text': "\n\n"
                }
            })
            requests.append({
                'insertInlineImage': {
                    'location': {'index': 1},
                    'uri': public_url,
                    'objectSize': {'width': {'magnitude': 500, 'unit': 'PT'}}
                }
            })

        docs_service.documents().batchUpdate(documentId=doc_id, body={'requests': requests}).execute()
        
        permission = {
            'type': 'user',
            'role': 'writer',
            'emailAddress': share_email
        }
        # We include supportsAllDrives=True to ensure this works in Shared Drive environments
        drive_service.permissions().create(
            fileId=doc_id, 
            body=permission, 
            fields='id',
            supportsAllDrives=True,
            sendNotificationEmail=True
        ).execute()
        
        doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"
        print(f"Successfully created and shared Google Doc: {doc_url}")
        return doc_url
        
    except HttpError as http_e:
        error_content = http_e.content.decode('utf-8')
        print(f"DEBUG: Google API Error Content: {error_content}")
        
        if http_e.resp.status == 403:
            return (
                f"Permission Denied (403): The current credentials do not have permission to access Google Docs or Drive. "
                f"Error Details: {error_content}\n\n"
                "COMMON FIXES:\n"
                "1. Ensure Docs and Drive APIs are enabled: https://console.cloud.google.com/apis/library/docs.googleapis.com?project=roitraining-dashboard\n"
                "2. Run: 'gcloud auth application-default login --scopes=https://www.googleapis.com/auth/cloud-platform,https://www.googleapis.com/auth/drive,https://www.googleapis.com/auth/documents'"
            )
        return f"Google API Error: {http_e} - {error_content}"
    except Exception as e:
        error_msg = f"Error creating Google Doc: {e}"
        print(error_msg)
        return error_msg

def save_text_report_to_gcs(company_name: str, report_text: str, bucket_name: str = "roitraining-dashboard-grounding") -> str:
    """
    Saves the raw text report to GCS as a .txt file.
    Naming: reports/Company_body_YYYYMMDD.txt
    """
    try:
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket_name)
        
        date_str = datetime.datetime.now().strftime("%Y%m%d")
        clean_company = company_name.replace(" ", "_")
        destination_path = f"reports/{clean_company}_body_{date_str}.txt"
        
        blob = bucket.blob(destination_path)
        blob.upload_from_string(report_text, content_type='text/plain')
        
        gcs_uri = f"gs://{bucket_name}/{destination_path}"
        print(f"Successfully uploaded raw text report to {gcs_uri}")
        return gcs_uri
        
    except Exception as e:
        error_msg = f"Error uploading text report to bucket: {e}"
        print(error_msg)
        return error_msg

def save_report_as_word(company_name: str, report_text: str, infographic_path: str) -> str:
    """
    Generates a Microsoft Word (.docx) report containing the infographic and analysis text.
    Layout: Infographic at the top, followed by Analysis text.
    Returns the local path of the generated Word document.
    """
    try:
        doc = docx.Document()
        
        # 1. Title
        title = doc.add_heading(f"TRIP Report - {company_name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2. Infographic
        if os.path.exists(infographic_path):
            doc.add_picture(infographic_path, width=Inches(6))
        else:
            doc.add_paragraph("[Infographic Placeholder]")
            
        doc.add_page_break()
        
        # 3. Report Text
        doc.add_heading("Analysis Report", level=1)
        
        sanitized_text = report_text.replace("’", "'").replace("“", '"').replace("”", '"').replace("–", "-")
        lines = sanitized_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.', line):
                doc.add_paragraph(line[line.find('.')+2:], style='List Number')
            else:
                doc.add_paragraph(line)

        date_str = datetime.datetime.now().strftime("%Y%m%d")
        clean_company = company_name.replace(" ", "_")
        local_docx = f"reports/{clean_company}_TRIP_Report_{date_str}.docx"
        
        doc.save(local_docx)
        print(f"Successfully generated Word document: {local_docx}")
        return local_docx
        
    except Exception as e:
        error_msg = f"Error generating Word document: {str(e)}"
        print(error_msg)
        return error_msg

def upload_file_to_drive(local_file_path: str, folder_id: str) -> str:
    """
    Uploads a local file to a designated Google Drive folder.
    Returns the file metadata or error message.
    """
    try:
        scopes = ['https://www.googleapis.com/auth/drive.file', 'https://www.googleapis.com/auth/drive']
        
        creds = None
        if not os.environ.get("K_SERVICE"):
            creds_path = "/Users/joey/Dev/sight_report/service-account-key.json"
        else:
            creds_path = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")
            
        if creds_path and os.path.exists(creds_path):
            creds = service_account.Credentials.from_service_account_file(creds_path, scopes=scopes)
        else:
            creds, _ = default(scopes=scopes)
            if hasattr(creds, 'with_scopes'):
                creds = creds.with_scopes(scopes)

        drive_service = build('drive', 'v3', credentials=creds)
        
        file_metadata = {
            'name': os.path.basename(local_file_path),
            'parents': [folder_id]
        }
        
        # Determine mime type based on extension
        ext = os.path.splitext(local_file_path)[1].lower()
        mime_type = 'application/octet-stream'
        if ext == '.docx':
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif ext == '.pdf':
            mime_type = 'application/pdf'
        elif ext == '.txt':
            mime_type = 'text/plain'
        elif ext in ['.png', '.jpg', '.jpeg']:
            mime_type = f'image/{ext[1:]}' if ext != '.jpg' else 'image/jpeg'

        media = MediaFileUpload(local_file_path, mimetype=mime_type, resumable=True)
        
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, webViewLink',
            supportsAllDrives=True
        ).execute()
        
        drive_link = file.get('webViewLink')
        print(f"Successfully uploaded {local_file_path} to Google Drive: {drive_link}")
        return drive_link

    except Exception as e:
        error_msg = f"Error uploading to Google Drive: {str(e)}"
        print(error_msg)
        return error_msg
